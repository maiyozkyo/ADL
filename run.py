import zipfile
from docx.oxml import OxmlElement
from docx import Document
from bs4 import BeautifulSoup
import xml.etree.ElementTree as ET
import huffman as hm




def setHiddenProperty(p):
    pPr = OxmlElement('w:pPr') # paragraph property
    rPr = OxmlElement('w:rPr') # run property
    v = OxmlElement('w:vanish') # hidden
    rPr.append(v)
    pPr.append(rPr)
    p._p.append(pPr)
    
def extractMsg(stegoDocxPath, extractMsgPath, tree):
    docZip = zipfile.ZipFile(stegoDocxPath)
    xmlStego = docZip.read('word/document.xml')
    xmlData = BeautifulSoup(xmlStego, 'xml').prettify()
    print(xmlData)
    root = ET.fromstring(xmlData)
    namespace = {'w': "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    
    textElements = root.findall(".//w:p", namespace)
    
    for ele in textElements:
        vanish = ele.findall('.//w:vanish', namespace)
        if (len(vanish) > 0):
            msg = ele.find('.//w:t', namespace).text.strip()
            if (msg != '' and msg != '\n'):
                extractMsg = hm.Huffman_Decoding(msg, tree)
                with open(extractMsgPath, 'w') as extractFile:
                    extractFile.write(extractMsg)
            

def hideMsg(coverDocxPath, msgDocxPath, stegoDocxPath):
    coverDoc = Document(coverDocxPath)
    p = coverDoc.add_paragraph()
    
    setHiddenProperty(p)  # set paragraph hidden property
    r = p.add_run()
    with open(msgDocxPath, 'r') as msgFile:
        msg = msgFile.read()
        huffmanMsg, tree = hm.Huffman_Encoding(msg)
        r.text = huffmanMsg
        r.font.hidden = True
        
    coverDoc.save(stegoDocxPath)
    return tree
    
coverFilePath = "test.docx"
msgDocxPath = "msgFile.txt"
# msgDocxPath = "msgFile.docx"
stegoXMLPath = "stego.xml"
stegoDocxPath = "stego.docx"
stegoDocxPath1 = "stego1.docx"
extractMsgPath = 'extractMsg.txt'

tree = hideMsg(coverFilePath, msgDocxPath, stegoDocxPath)
extractMsg(stegoDocxPath, extractMsgPath, tree)
